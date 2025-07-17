import streamlit as st
import base64
import fitz  # PyMuPDF
import docx
import os
import streamlit as st
from openai import OpenAI

openai_api_key = st.secrets["OPENAI_API_KEY"]
client = OpenAI(api_key=openai_api_key)

# --- Config ---
st.set_page_config(page_title="Estimating AI Dashboard", layout="wide")

# --- Password Gate ---
def check_password():
    def password_entered():
        if st.session_state["password"] == "Wbg3033!":
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.text_input("Enter password", type="password", on_change=password_entered, key="password")
        st.stop()
    elif not st.session_state["password_correct"]:
        st.error("❌ Incorrect password")
        st.stop()

check_password()

# --- Branding ---
def get_base64_image(image_path):
    with open(image_path, "rb") as img_file:
        encoded = base64.b64encode(img_file.read()).decode()
        return f"data:image/jpeg;base64,{encoded}"

encoded_logo = get_base64_image("logo.jpg")
st.markdown(
    f"""
    <div style='text-align: center;'>
        <img src="{encoded_logo}" style='width: 375px; margin-bottom: 10px;' />
        <h1 style='color: white; font-size: 36px;'>🧮 Estimating AI Dashboard</h1>
    </div>
    """,
    unsafe_allow_html=True
)

# --- Tabs Layout ---
tabs = st.tabs([
    "Estimate Analyzer", 
    "Takeoff Helper", 
    "Unit Pricing", 
    "Scope Review", 
    "Bid Summary", 
    "Proposal Checklist",
    "Proposal & Specs",
    "More Coming Soon"
])

with tabs[0]:
    st.subheader("📑 Estimate Analyzer")
    st.info("This can show basic project info (name, client, location, scope), estimator assigned, timeline (start/target completion). status, estimated cost.")
    st.info("Automation ideas: Autofill from forms (pdf, word, etc) and AI-generated short project summary based on scope")


with tabs[1]:
    st.subheader("📐 Takeoff Helper")
    st.info("Purpose: Track all material estimates and calculations")
    st.info("What to show: CSV upload or form input for material quantities, Auto-calculated material cost based on price database")
    st.info("Automation ideas: AI-powered OCR to pull takeoffs from drawings or PDFs, Price lookup from vendor APIs or price sheets, Auto-detect errors like unit mismatches or missing quantities")

with tabs[2]:
    st.subheader("💵 Unit Pricing Tool")
    st.info("Purpose: Calculate time and labor cost by task or phase")
    st.info("What to show: Labor hours per task, Cost per crew, Total labor cost")
    st.info("Automation ideas: Use AI to suggest labor hours based on past similar jobs, Pull crew rates from a central table, Predict labor bottlenecks or over-allocations")

with tabs[3]:
    st.subheader("📋 Scope Review Assistant")
    st.info("Purpose: Let AI analyze and optimize estimates")
    st.info("Ideas: Compare your numbers to similar past jobs and flag outlines, Suggest missing line items or hidden costs, Recommend markup based on market trends or workload")

with tabs[4]:
    st.subheader("📊 Bid Summary Generator")
    st.info("Purpose: Visual summary of where money is going")
    st.info("What to show: Pie chart or graph for: material, labor, equipment, subcontractors, overhead, profit")
    st.info("Automation ideas: Auto-generate charts from uploaded Excel/CSV files, Let estimators interactively adjust markup and see live effects")

    import pandas as pd
    import matplotlib.pyplot as plt

    cost_file = st.file_uploader("📥 Upload Bid Breakdown (CSV or Excel)", type=["csv", "xlsx"], key="bid_breakdown")

    if cost_file:
        try:
            if cost_file.name.endswith(".csv"):
                df = pd.read_csv(cost_file, header=None)
            else:
                df = pd.read_excel(cost_file, header=None)

            st.write("📋 Uploaded Raw Data Preview:")
            st.dataframe(df)

            # Find header row index for Material and Material Total
            header_row = None
            for idx, row in df.iterrows():
                row_lower = row.astype(str).str.lower()
                if "material" in row_lower.values and "material total" in row_lower.values:
                    header_row = idx
                    break

            if header_row is not None:
                # Re-read with correct header row
                if cost_file.name.endswith(".csv"):
                    df = pd.read_csv(cost_file, header=header_row)
                else:
                    df = pd.read_excel(cost_file, header=header_row)

                st.success(f"✅ Found headers at row {header_row + 1} — displaying material data.")
                st.dataframe(df)

                if "Material" in df.columns and "Material Total" in df.columns:
                    filtered_df = df[["Material", "Material Total"]].dropna()

                    st.write("📊 Material Cost Breakdown:")
                    st.dataframe(filtered_df)

                    fig, ax = plt.subplots()
                    ax.pie(filtered_df["Material Total"], labels=filtered_df["Material"], autopct="%1.1f%%", startangle=90)
                    ax.axis('equal')
                    st.pyplot(fig)
                else:
                    st.error("❌ Columns named **Material** and **Material Total** were not found after header detection.")
            else:
                st.error("❌ Could not find a header row with **Material** and **Material Total** in the same row.")

        except Exception as e:
            st.error(f"⚠️ Error processing file: {e}")

    # AI Summary Section
    def get_project_summary(text):
        if not text.strip():
            return "No text provided."

        response = client.chat.completions.create(
            model="gpt-4-1106-preview",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that summarizes construction project estimates."},
                {"role": "user", "content": f"Summarize this project: {text}"}
            ],
            temperature=0.3,
            max_tokens=300
        )
        return response.choices[0].message.content

    user_input = st.text_area("📝 Paste project text for AI summary", "", height=200)

    if st.button("🔍 Generate AI Summary"):
        summary = get_project_summary(user_input)
        st.success("✅ Summary Generated")
        st.text_area("🧠 AI Project Summary", summary, height=200)


with tabs[5]:
    st.subheader("✅ Proposal Checklist")

    uploaded_file = st.file_uploader("Upload a Proposal (PDF or Word)", type=["pdf", "docx"])

    def extract_text_from_pdf(file):
        text = ""
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
        return text

    def extract_text_from_docx(file):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            proposal_text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            proposal_text = extract_text_from_docx(uploaded_file)
        else:
            st.error("Unsupported file type.")
            st.stop()

        st.success("✅ Proposal uploaded and text extracted!")
        st.text_area("📄 Extracted Text", proposal_text, height=300)

        st.warning("⚠️ AI checklist comparison will be added once checklist/API is provided.")

with tabs[6]:
    st.subheader("↔️ Proposal & Specs Cross-Check")

    st.markdown("### 📤 Upload Proposal Document")
    proposal_file = st.file_uploader("Upload your proposal (PDF or DOCX)", type=["pdf", "docx"], key="proposal")

    st.markdown("### 📤 Upload Specification Document")
    specs_file = st.file_uploader("Upload specifications (PDF only)", type=["pdf"], key="specs")

    if proposal_file and specs_file:
        import fitz  # PyMuPDF
        from docx import Document

        def extract_text_from_pdf(file):
            doc = fitz.open(stream=file.read(), filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            return text

        def extract_text_from_docx(file):
            doc = Document(file)
            return "\n".join([p.text for p in doc.paragraphs])

        # Extract text
        if proposal_file.name.endswith(".pdf"):
            proposal_text = extract_text_from_pdf(proposal_file)
        else:
            proposal_text = extract_text_from_docx(proposal_file)

        specs_text = extract_text_from_pdf(specs_file)

       # Manual keyword check
        keywords = ["Thermal and Moisture Protection", "Metal Wall Panels", "ACM", "Metal Siding", "Soffit Panels", "Sheet Metal", "Sealants"]
        st.markdown("### 🔍 Manual Keyword Check in Specifications:")
        for keyword in keywords:
            if keyword.lower() in specs_text.lower():
                st.success(f"✅ Found '{keyword}' in specifications.")
            else:
                st.warning(f"⚠️ '{keyword}' not found in specifications.")

        # AI Cross-Check Function
        def get_ai_comparison(proposal_text, specs_text):
            prompt = f"""
You are a construction compliance assistant. Your job is to analyze two documents:
1. A project proposal
2. A specification document

Return a bullet list of any sections, materials, or language from the specification that seem to be missing or not clearly addressed in the proposal.

Proposal:
{proposal_text[:6000]}

Specifications:
{specs_text[:6000]}

Be specific and concise in your analysis.
"""
            response = client.chat.completions.create(
                model="gpt-4-1106-preview",
                messages=[
                    {"role": "system", "content": "You are a construction compliance assistant."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                max_tokens=500
            )
            return response.choices[0].message.content

        # Button to run AI check
        if st.button("🔎 Run AI Cross-Check"):
            comparison_result = get_ai_comparison(proposal_text, specs_text)
            st.success("✅ Cross-check complete!")
            st.markdown("### 📋 AI Findings:")
            st.markdown(comparison_result)


with tabs[7]:
    st.info("🚧 Stay tuned for more tools here!")
